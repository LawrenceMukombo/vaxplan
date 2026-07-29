from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\vaxplan")
SCREEN_DIR = ROOT / "comprehensive_user_guide_screenshots"
OUT = ROOT / "VaxPlan_Comprehensive_User_Guide.docx"

EXTRA = {
    "gis-point-intelligence-user": Path(
        r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-9b3f5c10-4701-4347-8afb-515d154ae60d.png"
    ),
    "facility-community-routes-user": Path(
        r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-76e424ae-e4ec-47ac-b570-344d9aa41dc7.png"
    ),
    "client-logbook-user": Path(
        r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-effa2ff4-d229-41ea-9321-9c067d405a99.png"
    ),
    "digital-vaccine-card-user": Path(
        r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-77ed9950-ff53-484d-ae6e-6a2719d96742.png"
    ),
}


def img(name: str):
    path = SCREEN_DIR / f"{name}.png"
    if path.exists():
        return path
    return EXTRA.get(name)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
styles["Normal"].font.size = Pt(10)
for style_name, size, color in [
    ("Heading 1", 16, (22, 82, 120)),
    ("Heading 2", 13, (31, 108, 146)),
    ("Heading 3", 11.5, (55, 96, 120)),
]:
    style = styles[style_name]
    style.font.name = "Aptos Display" if style_name == "Heading 1" else "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)


def h2(text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    r.font.size = Pt(10)


def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.22)
        r = p.add_run(item)
        r.font.size = Pt(9.7)


def steps(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(item)
        r.font.size = Pt(9.7)


def note(title, text, fill="EAF7FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(22, 82, 120)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(path, caption=None, width=6.65):
    if not path or not Path(path).exists():
        note("Screenshot pending", f"Missing screenshot file: {path}", fill="FFF3CD")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(8.4)
        r.font.color.rgb = RGBColor(91, 107, 117)
        cap.paragraph_format.space_after = Pt(8)


def module_table(rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(
        table.rows[0].cells,
        ["Module", "Primary users", "Main tasks", "Outputs / decisions"],
    ):
        cell_text(cell, text, bold=True, color=(255, 255, 255), size=8.5)
        shade(cell, "0B4F71")
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell_text(cell, text, size=8.2)
    doc.add_paragraph()


def cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VaxPlan Comprehensive User Guide")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = RGBColor(10, 43, 71)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Step-by-step operating manual for immunization microplanning, GIS catchment "
        "management, logistics, supervision, surveillance, reporting, and administration"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(62, 78, 88)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Zambia workspace example | Light theme screenshots | Generated from the running "
        "VaxPlan application and source-code review"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 108, 146)
    add_image(
        img("dashboard-overview"),
        "Dashboard example used as the operational starting point after sign-in.",
        width=6.8,
    )
    doc.add_page_break()


def write_intro():
    h1("How to Use This Guide")
    body(
        "This guide is written for programme managers, district and facility teams, "
        "monitoring and evaluation officers, ICT officers, logistics officers, supervisors, "
        "and partner users. It explains what each VaxPlan module is for and gives practical "
        "instructions for common daily tasks."
    )
    body(
        "Use the guide in three ways: as an onboarding manual for new users, as a refresher "
        "for experienced users, and as a reference during configuration, microplanning, "
        "supervision, and review meetings."
    )
    note(
        "Important",
        "Screenshots use the Republic of Zambia Ministry of Health workspace as an example. "
        "Menu options may vary depending on role, permissions, enabled modules, and active scope.",
    )
    h2("End-to-End Workflow at a Glance")
    module_table(
        [
            [
                "Setup and administration",
                "National admins, ICT, GIS officers",
                "Configure users, roles, boundaries, custom layers, facilities, staff, standards, integrations",
                "A secure, country-specific workspace ready for planning",
            ],
            [
                "Catchment preparation",
                "District teams, facility in-charges, GIS officers",
                "Review facilities, add communities, draw polygons, verify population, route and access evidence",
                "Validated facility catchments and service populations",
            ],
            [
                "Microplanning",
                "Facility teams, district planners, programme managers",
                "Build plans, forecast vaccines, plan sessions, staffing, budget, supervision",
                "Approved plan for implementation",
            ],
            [
                "Implementation",
                "Health workers, CHWs, logistics, supervisors",
                "Run sessions, track stock, follow defaulters, supervise field work, record readiness",
                "Evidence that planned services were delivered",
            ],
            [
                "Review and accountability",
                "Managers, M&E, partners",
                "Review reports, approvals, surveillance, missed communities, dropout and zero-dose indicators",
                "Decisions, corrective actions, and programme follow-up",
            ],
        ]
    )


def write_navigation_dashboard():
    h1("1. Sign-In, Navigation, and Common Controls")
    body(
        "After sign-in, VaxPlan opens into the dashboard with a persistent left navigation "
        "menu, active country selector, search, sync indicator, online presence, notifications, "
        "theme toggle, and user menu."
    )
    steps(
        [
            "Confirm the active country at the top of the page before editing records or creating plans.",
            "Use the left navigation groups to move between Main, Planning, SIA Campaigns, Analytics, Workflow, Administration, and System modules.",
            "Use the global search box to find features quickly.",
            "Check sync status before closing the browser, especially after offline work.",
            "Use the user menu for profile, settings, password, and logout actions.",
        ]
    )
    add_image(
        img("dashboard-overview"),
        "Dashboard and common app chrome: country context, sync state, global search, navigation, and user controls.",
    )
    h1("2. Dashboard")
    body(
        "The dashboard is the operational starting point. It summarizes the current planning "
        "period, facility context, equity signals, session and stock alerts, approvals, missed "
        "communities, and navigation shortcuts."
    )
    steps(
        [
            "Review the year, quarter, country, and facility scope at the top of the page.",
            "Scan equity cards for zero-dose, under-immunized children, dropout, missed communities, and denominator confidence.",
            "Open the relevant detail view from a dashboard card when a signal requires action.",
            "Use dashboard findings in weekly or monthly review meetings to agree follow-up actions.",
        ]
    )


def write_map_and_facilities():
    h1("3. GIS Map View and Point Intelligence")
    body(
        "The map is a planning workspace. It combines administrative boundaries, health facilities, "
        "communities, unserved places, session markers, custom layers, population overlays, and point intelligence."
    )
    h2("How to Review Planning Layers")
    steps(
        [
            "Open Map View from the Main menu.",
            "Confirm the active country and selected basemap.",
            "Use the legend to turn planning layers on or off.",
            "Zoom to the district or facility area before interpreting dense layers.",
            "Use Export only when authorized to share the map output.",
        ]
    )
    add_image(img("map-planning-layers"), "Map view showing planning layers, legend, facilities, sessions, and unserved places.")
    h2("How to Use Point Intelligence")
    steps(
        [
            "Open Map View and zoom to the area of interest.",
            "Click a location, boundary, marker, or planning point.",
            "Review coordinates, population estimates, nearby facility/community context, catchment proximity, and HTR status.",
            "Use quick actions such as zooming to the area or planning an outreach session when follow-up is required.",
        ]
    )
    add_image(EXTRA["gis-point-intelligence-user"], "Point intelligence popup: gridded population, nearby communities, catchment proximity, and HTR decision support.")

    h1("4. Facilities and Catchment Management")
    body(
        "Facilities are the anchor for routine immunization and campaign microplanning. A facility "
        "record stores basic details, geographic location, assigned communities, catchment polygon, "
        "staff roster, cold-chain assets, and population estimates."
    )
    add_image(img("facilities-registry"), "Facilities registry: facility records, population, assigned communities, staff, equipment, status, and edit actions.")
    h2("How to Add or Edit a Facility")
    steps(
        [
            "Open Facilities from the Main menu.",
            "Use Province and District filters to narrow the list.",
            "Select Add Facility for a new record or the pencil icon to edit an existing record.",
            "Complete facility name, HMIS code, facility type, district, latitude, longitude, staff count, cold-chain flag, power flag, and active status.",
            "Use Place Pin on the map if coordinates need to be updated visually.",
            "Select Update Facility & Catchment or Save to commit the changes.",
        ]
    )
    add_image(img("facility-process-01-general-catchment"), "Facility edit workspace: facility details, coordinate fields, status toggles, and map location tools.")
    note("Data quality check", "Only save facility coordinates after verifying the point is within the correct district and country boundary.")
    h2("How to Add Communities Served by a Facility")
    steps(
        [
            "Open the facility record and select Communities Served.",
            "Review communities already assigned to the facility.",
            "Select Register Community to add a new served community.",
            "Enter community name, coordinates, population where available, transport mode, access category, seasonality, and HTR status.",
            "Drop a pin or draw a polygon if geographic evidence is available.",
            "Save the community and confirm it appears under Communities Served.",
        ]
    )
    add_image(img("facility-process-02-communities-served"), "Communities Served tab: assigned communities, routes, distance, travel time, and register-community action.")
    add_image(img("facility-process-07-add-community-dialog"), "Add community workflow: community details, access characteristics, pin/polygon tools, and save action.")
    h2("How to Draw Facility or Community Polygons")
    steps(
        [
            "Open the facility edit dialog and select Polygon Drawing.",
            "Choose whether you are drawing the facility catchment boundary or a community sub-polygon.",
            "Select Draw Catchment or Draw Polygon Mode.",
            "Click around the intended boundary on the map. Use at least three points.",
            "Review the draft boundary visually. Clear and redraw if it crosses the wrong administrative boundary.",
            "Save the polygon when the boundary is verified.",
        ]
    )
    add_image(img("facility-process-03-polygon-drawing"), "Polygon Drawing tab: catchment and community polygon tools for map-based boundary maintenance.")
    h2("How to Add Staff Members and Community Health Workers")
    body(
        "The Staff Roster documents facility team members and community-based workers who support "
        "sessions, outreach, mobilization, supervision, logistics, and reporting."
    )
    steps(
        [
            "Open the facility record and select Staff Roster.",
            "Select Add Staff Member for facility-based staff or Add Community Worker where available.",
            "Enter name, role, phone number, email where available, cadre, assigned facility/community, and active status.",
            "Use consistent roles such as vaccinator, data clerk, supervisor, driver, cold-chain officer, community health worker, or mobilizer.",
            "Save the record and confirm the person appears in the roster before assigning them to session plans.",
        ]
    )
    add_image(img("facility-process-04-staff-roster"), "Staff Roster tab: facility staff and community worker management for session and outreach planning.")
    add_image(img("facility-process-08-add-staff-dialog"), "Add Staff Member dialog: registering a facility team member.")
    h2("How to Record Cold-Chain and Power Information")
    steps(
        [
            "Open the facility record and select Cold Chain.",
            "Review existing cold-chain equipment and readiness status.",
            "Record refrigerators, freezers, carriers, temperature monitoring, power availability, maintenance status, and alerts where applicable.",
            "Use summary Cold Chain and Power Supply toggles where only summary information is available.",
            "Save changes and use stock/logistics modules to follow up on equipment gaps.",
        ]
    )
    add_image(img("facility-process-05-cold-chain"), "Cold Chain tab: equipment, power, refrigerator, maintenance, and readiness information.")
    h2("How to Review Facility Population Evidence")
    steps(
        [
            "Open the facility record and select Population.",
            "Review facility-level population totals, catchment population, gridded population, and denominator assumptions.",
            "Compare facility estimates with assigned communities and settlement data.",
            "Use the result to support session planning, vaccine forecasting, and denominator confidence discussions.",
        ]
    )
    add_image(img("facility-process-06-population"), "Facility population tab: denominator and catchment population evidence used in planning.")
    add_image(EXTRA["facility-community-routes-user"], "Facility/community route evidence showing distance, walking time, seasonality, and referral paths.")


def write_settlements_microplans():
    h1("5. Settlements, Population Hub, and Denominator Management")
    body(
        "Settlements and population data help teams identify where people live, estimate service needs, "
        "and detect locations that may not be covered by routine sessions."
    )
    h2("Settlements Registry")
    steps(
        [
            "Open Settlements from the Main menu.",
            "Filter by province, district, facility, risk, or status.",
            "Review settlement name, location, nearest facility, population, distance, travel time, linked community, and risk.",
            "Use settlement evidence to update facility catchments or investigate missed/unserved places.",
        ]
    )
    add_image(img("settlements-registry"), "Settlements registry: community-level planning records, distances, risk, and catchment linkage.")
    h2("Population Hub")
    steps(
        [
            "Open Population Hub.",
            "Select year, province, district, or facility scope.",
            "Review imported grids, administrative estimates, facility catchment totals, and assumptions.",
            "Use population outputs in microplanning, vaccine forecasting, coverage review, and gap analysis.",
        ]
    )
    add_image(img("population-hub"), "Population Hub: denominator and population planning workspace.")

    h1("6. Routine and Campaign Microplanning")
    body(
        "Microplanning is a bottom-up process that starts at facility catchment level. VaxPlan "
        "structures microplanning into a 12-step workflow covering coverage review, communities, "
        "risk, sessions, staffing, vaccines, demand generation, logistics, budget, supervision, "
        "approval, and execution review."
    )
    steps(
        [
            "Open Routine Microplan or SIA Campaign Microplan.",
            "Create a new plan or open an existing draft.",
            "Move through the steps from left to right, saving as draft when needed.",
            "Resolve validation warnings before submission.",
            "Submit the plan for review when complete.",
        ]
    )
    add_image(img("routine-microplan-list"), "Routine microplan list: draft, submitted, reviewed, and approved plans.")
    wizard_steps = [
        ("Step 1. Coverage and denominators", "Review prior coverage, dropout, target population, stockouts, AEFI, and denominator assumptions.", "microplan-wizard-step-01"),
        ("Step 2. Catchment and communities", "Confirm assigned communities, population targets, catchment gaps, and community coverage.", "microplan-wizard-step-02"),
        ("Step 3. Risk scoring", "Score hard-to-reach and service-risk factors to prioritize communities needing special attention.", "microplan-wizard-step-03"),
        ("Step 4. Session calendar", "Plan fixed, outreach, mobile, or special sessions with dates and community targets.", "microplan-wizard-step-04"),
        ("Step 5. Staffing per session day", "Assign staff, supervisors, team type, daily target, and per-diem assumptions.", "microplan-wizard-step-05"),
        ("Step 6. Vaccine forecasting", "Calculate antigen requirements, wastage, vials, syringes, safety boxes, and cold-chain needs.", "microplan-wizard-step-06"),
        ("Step 7. Demand generation", "Plan mobilization channels, focal persons, community volunteer work, and HFC readiness.", "microplan-wizard-step-07"),
        ("Step 8. Logistics and transport", "Record transport mode, route assumptions, distance, fuel, access barriers, and movement support.", "microplan-wizard-step-08"),
        ("Step 9. Budget", "Capture cost items, funding gaps, source of funds, and approval needs.", "microplan-wizard-step-09"),
        ("Step 10. Supervision plan", "Plan supervision visits, supervisors, checklists, dates, and follow-up responsibilities.", "microplan-wizard-step-10"),
        ("Step 11. Submit for approval", "Review completeness checks, validation messages, and submit the microplan for approval.", "microplan-wizard-step-11"),
        ("Step 12. Execution and review", "Track implementation, completed sessions, actuals, gaps, and lessons for the next planning cycle.", "microplan-wizard-step-12"),
    ]
    for title, desc, shot in wizard_steps:
        h2(title)
        body(desc)
        add_image(img(shot), f"{title}: screenshot from a populated Zambia routine microplan.", width=6.55)


def write_operations_clients_supervision():
    h1("7. Sessions, Stock, Readiness, and Hard-to-Reach Follow-up")
    h2("Sessions Hub")
    body(
        "The Sessions Hub consolidates planned, in-progress, completed, overdue, and historical "
        "immunization sessions."
    )
    steps(
        [
            "Open Sessions.",
            "Filter by year, quarter, province, district, facility, status, or session type.",
            "Open a session to review planned communities, team, date, status, and outputs.",
            "Update progress or follow up overdue sessions according to your role.",
        ]
    )
    add_image(img("sessions-hub"), "Sessions Hub: consolidated session planning and implementation tracking.")
    h2("Stock Ledger and Vaccine Logistics")
    steps(
        [
            "Open Stock Ledger.",
            "Select facility, antigen, period, or stock status filters.",
            "Review opening balance, receipts, issues, adjustments, losses, and closing balance.",
            "Compare stock levels with upcoming session plans and vaccine forecasts.",
            "Escalate risks before planned sessions are affected.",
        ]
    )
    add_image(img("stock-ledger"), "Stock Ledger: vaccine and logistics visibility linked to readiness.")
    h2("Plan Health, Field Readiness, Hard-to-Reach, Missed Communities, and Dropout")
    body("These modules help users detect and correct planning weaknesses before service delivery.")
    for shot, caption in [
        ("plan-health", "Plan Health: completeness and readiness checks for microplans."),
        ("field-readiness", "Field Readiness: operational readiness signals for field implementation."),
        ("hard-to-reach", "Hard-to-Reach: risk scoring and prioritization."),
        ("missed-communities", "Missed Communities: unserved or missed places requiring planning follow-up."),
        ("dropout-rates", "Dropout Rates: DTP/MCV and other dropout indicators for follow-up."),
    ]:
        add_image(img(shot), caption)

    h1("8. Client Logbook, Defaulter Follow-up, and Digital Vaccine Card")
    body(
        "The Client Logbook links catchment microplanning to individual service follow-up. It helps "
        "teams find children or pregnant women due for services, send reminders, identify defaulters, "
        "and review vaccination status."
    )
    h2("How to Use the Client Registry")
    steps(
        [
            "Open Client Logbook.",
            "Select Active Registry, Cohort Due Queue, or Defaulters.",
            "Filter by province, district, facility, date, cohort, or status.",
            "Search by client name, parent or guardian, phone number, or village.",
            "Use SMS or follow-up actions according to national data protection rules.",
        ]
    )
    add_image(img("client-logbook-registry"), "Client Logbook: active registry, due queue, defaulters, filters, and actions.")
    add_image(EXTRA["client-logbook-user"], "Client registry example: due queues, defaulter counts, SMS reminders, and vaccination-card access.")
    h2("How to Review a Digital Vaccine Card")
    steps(
        [
            "Open Client Logbook and find the client.",
            "Select Vax Card or the vaccination-card action.",
            "Review client identity, date of birth, guardian, contact, registered clinic, and catchment.",
            "Review each antigen/dose row for target age, status, given date, clinic, batch number, and VVM.",
            "Use missed and pending dose warnings to plan defaulter tracing or the next visit.",
            "Do not administer doses marked clinically ineligible; follow national schedule rules.",
        ]
    )
    add_image(img("client-digital-vaccine-card"), "Digital vaccine card: dose grid, missed doses, pending doses, batch/VVM information, and catchment identity.")
    add_image(EXTRA["digital-vaccine-card-user"], "Detailed digital vaccine card example with missed-dose and catch-up eligibility warning information.")
    add_image(img("defaulter-list"), "Defaulter List: clients requiring tracing, reminders, and catch-up planning.")

    h1("9. Supportive Supervision, Checklist Builder, PCE, and House-to-House Monitoring")
    body(
        "Supportive supervision helps managers verify service quality, readiness, documentation, and "
        "corrective actions. VaxPlan supports routine supervision, campaign supervision tools, checklist "
        "templates, post-campaign evaluation, and house-to-house monitoring."
    )
    h2("How to Plan and Record Supportive Supervision")
    steps(
        [
            "Open Supervision from the System menu.",
            "Review scheduled, completed, and pending supervision activities.",
            "Create or open a supervision visit.",
            "Select facility, campaign or routine context, date, supervisor, checklist, and focus area.",
            "Complete checklist responses during or after the visit.",
            "Record findings, actions, responsible person, due date, and status.",
            "Use reports and dashboards to follow unresolved actions.",
        ]
    )
    add_image(img("supportive-supervision-dashboard"), "Supportive Supervision dashboard: visits, checklists, findings, and action tracking.")
    h2("How to Develop a Supportive Supervision Checklist")
    steps(
        [
            "Open Supervision Templates.",
            "Select New Template or Create Template.",
            "Give the checklist a clear name, purpose, programme area, and applicable level.",
            "Add sections such as cold chain, vaccine stock, documentation, AEFI, waste management, session readiness, data quality, and community mobilization.",
            "Add questions using yes/no, score, text, number, date, select option, or evidence upload where supported.",
            "Mark critical questions that require follow-up when failed.",
            "Save or publish the template after review.",
            "Use the template in a supervision visit and revise it based on field feedback.",
        ]
    )
    add_image(img("supervision-template-list"), "Supervision templates: checklist template library.")
    add_image(img("supervision-template-builder"), "Checklist builder: create structured supervision templates with sections and questions.")
    h2("Campaign Supervision, PCE, and House-to-House Monitoring")
    body("Campaign workflows focus on SIA readiness, in-process monitoring, post-campaign evaluation, and household-level verification.")
    for shot, caption in [
        ("campaign-supervision-tools", "Campaign Supervision Tools: campaign checklist and monitoring workspace."),
        ("post-campaign-evaluation", "Post-Campaign Evaluation: campaign review and performance assessment workspace."),
        ("house-to-house-monitoring", "House-to-House Monitoring: household monitoring and missed-child follow-up evidence."),
    ]:
        add_image(img(shot), caption)


def write_system_admin_sop():
    h1("10. VPD Surveillance, Reports, Approvals, and Notifications")
    h2("VPD Surveillance")
    steps(
        [
            "Open Surveillance.",
            "Use Dashboard for surveillance indicators and trends.",
            "Use Case Linelist to search, filter, and open case records.",
            "Use Spatial View to map cases and reporting facilities.",
            "Use Report New Case to enter disease, classification, patient, onset date, facility, coordinates, investigation date, and clinical notes.",
            "Use Configuration where authorized to maintain linelist templates and alert routing.",
        ]
    )
    add_image(img("vpd-surveillance-dashboard"), "VPD Surveillance dashboard: surveillance metrics, disease trends, and report-new-case action.")
    h2("Reports and Dashboards")
    steps(
        [
            "Open Reports.",
            "Select reporting year, quarter, province, district, facility, and active location scope.",
            "Choose the report tab such as sessions, microplans, zero-dose, missed communities, coverage, hard-to-reach, budget, or supervision.",
            "Review charts and tables before export.",
            "Use Export only for authorized reporting and partner review.",
        ]
    )
    add_image(img("reports-dashboard"), "Reports dashboard: structured programme review outputs and filters.")
    h2("Approvals and Notifications")
    body("Approval workflows support accountability. Notifications keep users aware of submissions, returned plans, assigned actions, system alerts, and follow-up requirements.")
    add_image(img("approvals-workflow"), "Approvals workflow: review, approve, return, or track submitted plans and requests.")
    add_image(img("notifications-center"), "Notifications center: workflow follow-up, alerts, and system messages.")

    h1("11. Administration, RBAC, Staff, Boundaries, Layers, and Integrations")
    h2("User Management and Role-Based Access")
    steps(
        [
            "Open User Management.",
            "Search for the user by name or email.",
            "Review role, facility, district, province, status, and tenant assignment.",
            "Assign the minimum role needed for the user to perform their duties.",
            "Deactivate users who no longer need access.",
            "Use Access Requests to approve or reject new signup requests.",
        ]
    )
    add_image(img("user-management-rbac"), "User Management: accounts, roles, scopes, and access controls.")
    add_image(img("access-requests"), "Access Requests: review and approve pending user registration requests.")
    h2("Manage Staff")
    body("The Manage Staff module provides a broader administrative view of personnel records, while the facility Staff Roster manages staff attached to a specific facility.")
    add_image(img("staff-management"), "Manage Staff: administrative staff registry for personnel and role maintenance.")
    h2("Boundaries, Custom Layers, and HIS Integrations")
    steps(
        [
            "Use Boundary Manager to upload or maintain administrative boundaries such as province, district, ward, or catchment GeoJSON.",
            "Use Custom Layers to add operational layers such as gridded population, points of interest, education facilities, roads, or programme-specific overlays.",
            "Use HIS Integrations to configure external system linkages such as DHIS2 or interoperability endpoints.",
            "Validate imported data before using it for official planning or reporting.",
        ]
    )
    add_image(img("boundary-manager"), "Boundary Manager: administrative and planning boundary management.")
    add_image(img("custom-layers"), "Custom Layers: upload and configure geospatial overlays.")
    add_image(img("his-integrations"), "HIS Integrations: interoperability and connected system configuration.")
    h2("National Plan, Data Sources, Standards, Settings, and Help")
    body("These modules support governance, transparency, configuration, and user support.")
    for shot, caption in [
        ("national-plan", "National Plan: annual planning objectives and programme planning context."),
        ("data-sources", "Data Sources: acknowledgements and source transparency."),
        ("standards-alignment", "Standards Alignment: WHO, UNICEF, Gavi, and national programme alignment evidence."),
        ("settings", "Settings: profile, preferences, security, and application configuration."),
        ("help-user-support", "Help and User Support: practical help, FAQs, and user guidance."),
    ]:
        add_image(img(shot), caption)

    h1("12. Practical Operating Procedures")
    h2("Before Creating a Microplan")
    bullets(
        [
            "Confirm facility coordinates and catchment boundary are correct.",
            "Confirm communities served are assigned and HTR status is accurate.",
            "Review denominator evidence in Population Hub and facility Population tab.",
            "Check staff roster and community worker availability.",
            "Check cold-chain readiness and stock risks.",
            "Review missed communities, dropout, and defaulter lists for follow-up priorities.",
        ]
    )
    h2("Before Submitting a Microplan")
    bullets(
        [
            "All 12 wizard steps have been reviewed.",
            "Session dates, sites, and communities are realistic.",
            "Vaccine forecast and logistics match planned sessions.",
            "Budget gaps are documented.",
            "Supervision plan is assigned to named supervisors.",
            "Validation warnings have been resolved or explained.",
        ]
    )
    h2("Before a Supervision Visit")
    bullets(
        [
            "Use the correct checklist template for the visit type.",
            "Confirm facility, date, supervisor, and programme focus.",
            "Review previous findings and pending actions.",
            "Carry or prepare evidence required for checklist questions.",
            "Record findings and actions before closing the visit.",
        ]
    )
    h2("Data Protection and Quality")
    bullets(
        [
            "Only access client records for authorized programme work.",
            "Do not export identifiable client data unless authorized.",
            "Verify coordinates and boundaries before saving.",
            "Use consistent facility, community, staff, and antigen names.",
            "Do not overwrite verified records without a reason.",
            "Report role or permission problems to the system administrator.",
        ]
    )


cover()
write_intro()
write_navigation_dashboard()
write_map_and_facilities()
write_settlements_microplans()
write_operations_clients_supervision()
write_system_admin_sop()

props = doc.core_properties
props.title = "VaxPlan Comprehensive User Guide"
props.subject = "User manual for VaxPlan health microplanning platform"
props.author = "VaxPlan / Codex documentation build"
doc.save(OUT)
print(OUT)
print(OUT.stat().st_size)
