from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

base = Path(r'C:\vaxplan\VaxPlan_Consolidated_Technical_Documentation.docx')
out = Path(r'C:\vaxplan\VaxPlan_Consolidated_Technical_Documentation_With_Temporal_Framework.docx')
diag = Path(r'C:\vaxplan\Diagrams\temporal')

doc = Document(str(base))
styles = doc.styles
for name, size in [('Normal', 10), ('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
    if name in styles:
        styles[name].font.name = 'Arial'
        styles[name].font.size = Pt(size)
        if name.startswith('Heading'):
            styles[name].font.color.rgb = RGBColor(31, 78, 121)

def p(text='', style=None, bold=False, color=None):
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
        hdr[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w'))))
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = tbl.add_row().cells
        for i,val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in tbl.rows:
            for idx,width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return tbl

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc.add_page_break()
p('Enterprise Time Dimension and Historical Reference Data Framework', 'Heading 1')
p('Implementation status: additive framework implemented in source code and database migration on 2026-07-19.', bold=True)
p('This section documents the bitemporal architecture added to VaxPlan so that ministries can manage current, historical, future-dated, retroactive, corrected, superseded, and cancelled reference data without destructive overwrites.')

p('Why This Matters', 'Heading 2')
bullets([
    'Immunization planning depends on facts that change over time: facility names, catchment boundaries, staff postings, role assignments, population denominators, and administrative hierarchy.',
    'Historical reports must remain explainable even after a district boundary, staff assignment, or denominator is corrected.',
    'Future-dated changes, such as a facility opening or boundary update, need approval before they affect maps, reports, and planning workflows.',
])

p('Bitemporal Design', 'Heading 2')
p('The framework separates real-world effective time from system-recorded time. Valid time answers what was true for operations. System time answers when VaxPlan recorded or trusted that fact.')
table(['Concept','Fields','Meaning'],[
    ['Valid time','valid_from, valid_to','The real-world period when the version is effective.'],
    ['System time','recorded_at, recorded_until','The period when the version was recorded and trusted in VaxPlan.'],
    ['Version state','status, is_current, is_future, is_correction','Workflow and lifecycle state.'],
    ['Governance','change_type, reason, summary, source, actors','Why the change happened, who made it, and what evidence supports it.'],
    ['Snapshot','snapshot, affected_records, metadata','The preserved record image and impact context.'],
],[1.3,1.8,3.6])

p('Temporal Data Flow', 'Heading 2')
doc.add_picture(str(diag/'vaxplan-temporal-data-flow.png'), width=Inches(6.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Figure: Temporal changes begin as drafts, pass overlap and impact assessment, move through approval, and then become active, scheduled, rejected, or preserved as audit evidence.')

p('Core Schema Added', 'Heading 2')
table(['Table','Purpose'],[
    ['temporal_entity_versions','Generic bitemporal snapshots for governed records including facilities, settlements, boundaries, reference data, and configuration.'],
    ['temporal_change_requests','Approval workflow for submitted, retroactive, future-dated, corrected, and rejected changes.'],
    ['temporal_audit_events','Immutable audit trail for creation, submission, approval, correction, cancellation, and rejection.'],
    ['temporal_entity_lineage','Split, merge, transfer, reparenting, replacement, and other lineage relationships.'],
    ['temporal_role_assignments','Effective user role, permission, and data-scope history.'],
    ['temporal_employment_assignments','Staff employment, facility posting, and administrative placement history.'],
    ['temporal_geography_versions','Administrative geography, hierarchy, code, name, and geometry versions.'],
    ['temporal_population_denominators','Reference-year population estimates, approved planning values, source confidence, and effective periods.'],
],[2.2,4.6])

p('Schema Overview', 'Heading 2')
doc.add_picture(str(diag/'vaxplan-temporal-schema-overview.png'), width=Inches(6.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Figure: Temporal governance tables are additive and link to the existing tenant, user, facility, geography, population, and planning modules through stable entity identifiers and source record ids.')

p('Lifecycle and State Model', 'Heading 2')
doc.add_picture(str(diag/'vaxplan-bitemporal-state-model.png'), width=Inches(6.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Figure: Records are never deleted to hide history. They move through draft, pending approval, active, scheduled, superseded, corrected, cancelled, or rejected states.')

p('API and User Interface', 'Heading 2')
p('The backend exposes secured endpoints under /api/temporal for inventory, current version lookup, as-of reconstruction, history, future changes, comparison, creation, submission, approval, rejection, correction, and cancellation. The frontend adds a Temporal History Workbench at /temporal-history and a sidebar entry under System.')
table(['Workflow','Endpoint family','Primary permission'],[
    ['Current review','GET /api/temporal/:entityType/:entityId/current','temporal.view'],
    ['History and as-of review','GET /history, /as-of, /compare','temporal.view_history'],
    ['Draft proposal','POST /versions','temporal.propose_change'],
    ['Approval decision','POST /approve, /reject','temporal.approve_change or temporal.review_change'],
    ['Correction and cancellation','POST /correct, /cancel','temporal.correct_history or temporal.cancel_future_change'],
],[1.8,2.9,2.0])

p('Role and Permission Defaults', 'Heading 2')
bullets([
    'National Admin: full temporal governance including approval, correction, cancellation, export, full audit, and configuration.',
    'National Manager: temporal view, history, and review access.',
    'GIS Specialist: temporal view, history, and proposed-change access for spatial and reference updates.',
])

p('Migration and Backfill Approach', 'Heading 2')
bullets([
    'Apply migration 0015_enterprise_temporal_framework.sql using the existing migration runner.',
    'Backfill high-priority current records by tenant using source record ids: facilities, geography, users, roles, staff assignments, settlements, catchments, and population denominators.',
    'Mark backfilled records with source_system = legacy_migration and source_type = system_backfill.',
    'Do not delete, rewrite, or reset existing operational records during temporal adoption.',
    'Validate duplicate current versions, overlapping periods, missing tenant ids, and orphan source ids before opening the workbench to production users.',
])

p('Implementation References', 'Heading 2')
bullets([
    'shared/schema.ts defines the temporal tables and insert/select types.',
    'migrations/0015_enterprise_temporal_framework.sql creates tables, constraints, indexes, and partial uniqueness for current versions.',
    'server/services/temporalService.ts implements bitemporal queries, overlap detection, approvals, corrections, cancellations, and audit events.',
    'server/routes/temporal.ts registers secured temporal API endpoints.',
    'client/src/pages/TemporalRecords.tsx implements the Temporal History Workbench.',
])

p('Verification', 'Heading 2')
p('The implementation passed TypeScript verification with npm run check. The framework is additive, so existing operational flows remain intact while the new temporal API and workbench are introduced.')

doc.save(str(out))
print(out)
