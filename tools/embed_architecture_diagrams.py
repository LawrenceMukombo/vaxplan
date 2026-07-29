from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\vaxplan")
SOURCE_DOC = ROOT / "VaxPlan_Consolidated_Technical_Documentation.docx"
OUTPUT_DOC = ROOT / "VaxPlan_Consolidated_Technical_Documentation_With_Architecture_Diagrams.docx"
DIAGRAM_DIR = ROOT / "Diagrams" / "light-theme-architecture"

NAVY = "#0B1F3A"
BLUE = "#1186B2"
TEAL = "#16A085"
CYAN = "#E8F7FC"
MINT = "#EAFBF5"
AMBER = "#FFF4DC"
ROSE = "#FFF1F2"
LILAC = "#F3F0FF"
GRAY = "#667085"
LIGHT = "#F8FAFC"
BORDER = "#CBD5E1"
DARK = "#1F2937"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = font(34, True)
FONT_SUBTITLE = font(18)
FONT_HEADER = font(21, True)
FONT_BODY = font(16)
FONT_SMALL = font(13)
FONT_TINY = font(11)


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline=BORDER, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_lines(draw, text, x, y, max_width, fnt, fill=DARK, line_gap=5, align="left"):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        tx = x
        if align == "center":
            tx = x + (max_width - (bbox[2] - bbox[0])) / 2
        draw.text((tx, y), line, fill=fill, font=fnt)
        y += bbox[3] - bbox[1] + line_gap
    return y


def label_box(draw, box, title, body, fill=WHITE, accent=BLUE):
    x1, y1, x2, y2 = box
    rounded(draw, box, fill=fill, outline=BORDER)
    draw.rounded_rectangle((x1, y1, x1 + 9, y2), radius=8, fill=accent)
    draw.text((x1 + 22, y1 + 16), title, fill=NAVY, font=FONT_HEADER)
    text_lines(draw, body, x1 + 22, y1 + 49, x2 - x1 - 44, FONT_BODY, fill=GRAY, line_gap=4)


def arrow(draw, start, end, color=BLUE, width=4):
    draw.line((start, end), fill=color, width=width)
    sx, sy = start
    ex, ey = end
    import math

    angle = math.atan2(ey - sy, ex - sx)
    length = 16
    spread = 0.55
    pts = [
        (ex, ey),
        (ex - length * math.cos(angle - spread), ey - length * math.sin(angle - spread)),
        (ex - length * math.cos(angle + spread), ey - length * math.sin(angle + spread)),
    ]
    draw.polygon(pts, fill=color)


def canvas(title, subtitle):
    img = Image.new("RGB", (1800, 1100), LIGHT)
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 96), fill=WHITE)
    draw.rectangle((0, 94, 1800, 96), fill="#DCEAF3")
    draw.rounded_rectangle((52, 24, 96, 68), radius=12, fill=BLUE)
    draw.text((112, 23), title, fill=NAVY, font=FONT_TITLE)
    draw.text((113, 61), subtitle, fill=GRAY, font=FONT_SUBTITLE)
    draw.text((1565, 35), "VaxPlan", fill=BLUE, font=FONT_HEADER)
    return img, draw


def save(img, name):
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    img.save(path, "PNG")
    return path


def system_context():
    img, draw = canvas("System Context Architecture", "How users, partners, and external systems interact with VaxPlan")
    center = (620, 270, 1180, 820)
    rounded(draw, center, fill=WHITE, outline=BLUE, radius=24, width=4)
    draw.text((770, 302), "VaxPlan Platform", fill=NAVY, font=FONT_TITLE)
    text_lines(
        draw,
        "GIS-enabled immunization microplanning, session planning, stock visibility, supervision, surveillance, reporting, and accountability.",
        705,
        355,
        390,
        FONT_BODY,
        fill=GRAY,
        align="center",
    )
    label_box(draw, (735, 470, 1065, 585), "Web / PWA App", "React interface for national, district, facility, GIS, M&E, and partner users.", CYAN, BLUE)
    label_box(draw, (735, 625, 1065, 740), "API + Data Services", "Tenant-aware services enforce permissions and coordinate workflows.", MINT, TEAL)

    actors = [
        ((90, 180, 420, 310), "MoH & EPI Teams", "National and subnational planning, approval, supervision, and reporting.", BLUE),
        ((90, 440, 420, 570), "Facility Teams", "Catchment review, community registration, sessions, stock, and field readiness.", TEAL),
        ((90, 700, 420, 830), "Partners", "Programme support, dashboards, review outputs, and implementation coordination.", "#7C3AED"),
        ((1380, 180, 1710, 310), "GIS / Population Data", "Boundaries, settlements, gridded population, routes, and POIs.", "#F97316"),
        ((1380, 440, 1710, 570), "Health Systems", "DHIS2, stock systems, IDSR, SMS, identity providers, and exports.", "#DC2626"),
        ((1380, 700, 1710, 830), "Offline Devices", "Field use with local cache, outbox replay, and sync conflict review.", "#0EA5E9"),
    ]
    for box, title, body, accent in actors:
        label_box(draw, box, title, body, WHITE, accent)
        if box[0] < 500:
            arrow(draw, (box[2], (box[1] + box[3]) // 2), (center[0], (box[1] + box[3]) // 2), color=accent)
        else:
            arrow(draw, (center[2], (box[1] + box[3]) // 2), (box[0], (box[1] + box[3]) // 2), color=accent)
    draw.text((95, 990), "Architecture principle: one national platform, tenant isolation, role-scoped workflows, and auditable data movement.", fill=GRAY, font=FONT_BODY)
    return save(img, "01-system-context-light.png")


def runtime_architecture():
    img, draw = canvas("Runtime and Deployment Architecture", "Application runtime components from browser to database")
    boxes = [
        ((70, 210, 390, 360), "User Browser / PWA", "React + Vite client, service worker, cache, local storage, installable PWA.", CYAN, BLUE),
        ((510, 210, 830, 360), "Express API Server", "REST routes, auth middleware, RBAC checks, validation, audit logging.", MINT, TEAL),
        ((950, 210, 1270, 360), "Service Modules", "Microplanning, GIS, stock, surveillance, supervision, reports, notifications.", LILAC, "#7C3AED"),
        ((1390, 210, 1710, 360), "PostgreSQL Database", "Drizzle schema tables, tenant-scoped records, planning and GIS entities.", AMBER, "#D97706"),
        ((510, 540, 830, 690), "Background Jobs", "Population refresh, import processing, sync replay, notifications, exports.", WHITE, "#0EA5E9"),
        ((950, 540, 1270, 690), "File / Import Assets", "CSV imports, custom layers, docs, GeoJSON, boundary and raster references.", WHITE, "#64748B"),
        ((1390, 540, 1710, 690), "External Services", "Map tiles, SMS, OIDC, DHIS2 or national reporting integrations.", ROSE, "#E11D48"),
    ]
    for box, title, body, fill, accent in boxes:
        label_box(draw, box, title, body, fill, accent)
    arrow(draw, (390, 285), (510, 285), BLUE)
    arrow(draw, (830, 285), (950, 285), TEAL)
    arrow(draw, (1270, 285), (1390, 285), "#D97706")
    arrow(draw, (670, 360), (670, 540), "#0EA5E9")
    arrow(draw, (1110, 360), (1110, 540), "#64748B")
    arrow(draw, (1270, 615), (1390, 615), "#E11D48")
    arrow(draw, (1390, 315), (1270, 315), "#94A3B8", width=3)
    draw.text((75, 870), "Deployment notes", fill=NAVY, font=FONT_HEADER)
    bullets = [
        "The browser client communicates with same-origin API endpoints using authenticated sessions.",
        "The API layer is the security boundary: tenant, role, permission, and geographic scope checks happen server-side.",
        "Operational modules share the same relational data model so maps, plans, stock, supervision, and reports stay connected.",
    ]
    y = 910
    for b in bullets:
        draw.ellipse((78, y + 5, 88, y + 15), fill=BLUE)
        y = text_lines(draw, b, 105, y, 1520, FONT_BODY, fill=GRAY, line_gap=3) + 8
    return save(img, "02-runtime-deployment-light.png")


def offline_sync():
    img, draw = canvas("Offline-First PWA and Synchronization", "How field changes move between browser cache and the server")
    stages = [
        ((90, 250, 390, 420), "1. Load Baseline", "User opens module. App retrieves server data and stores usable offline state.", BLUE),
        ((510, 250, 810, 420), "2. Work Offline", "Facility staff can update records, sessions, or field forms while connectivity is weak.", TEAL),
        ((930, 250, 1230, 420), "3. Queue Outbox", "Write actions are held locally with metadata, target endpoint, and user context.", "#F97316"),
        ((1350, 250, 1650, 420), "4. Replay to API", "When online, queued changes are sent to the authenticated API for validation.", "#7C3AED"),
        ((930, 620, 1230, 790), "5. Conflict Review", "Rejected or conflicting edits are surfaced for correction rather than silently overwritten.", "#DC2626"),
        ((510, 620, 810, 790), "6. Refresh Cache", "Accepted server state updates the local cache, dashboards, maps, and reports.", "#0EA5E9"),
    ]
    for box, title, body, accent in stages:
        label_box(draw, box, title, body, WHITE, accent)
    arrow(draw, (390, 335), (510, 335), BLUE)
    arrow(draw, (810, 335), (930, 335), TEAL)
    arrow(draw, (1230, 335), (1350, 335), "#F97316")
    arrow(draw, (1500, 420), (1080, 620), "#DC2626")
    arrow(draw, (930, 705), (810, 705), "#0EA5E9")
    arrow(draw, (660, 620), (240, 420), "#94A3B8", width=3)
    draw.text((95, 900), "Key controls: authenticated replay, server-side validation, audit logs, sync status, conflict screen, and user-visible retry behavior.", fill=GRAY, font=FONT_BODY)
    return save(img, "03-offline-sync-light.png")


def rbac_security():
    img, draw = canvas("Authentication, RBAC, and Geographic Scope", "Security decision path for protected actions")
    nodes = [
        ((95, 230, 395, 360), "User Identity", "Password/OIDC/local-dev login establishes a session and tenant context.", BLUE),
        ((510, 230, 810, 360), "Session Middleware", "Checks authentication, expiry, idle timeout, and active tenant.", TEAL),
        ((925, 230, 1225, 360), "Effective Permissions", "Combines role defaults with tenant custom roles and cached permissions.", "#7C3AED"),
        ((1340, 230, 1640, 360), "Action Authorization", "Endpoint validates permission for create, update, delete, approve, export, or view.", "#F97316"),
        ((925, 570, 1225, 700), "Geographic Scope", "Facility, district, province, tenant, or platform scope filters accessible records.", "#0EA5E9"),
        ((510, 570, 810, 700), "Audit Trail", "Writes are recorded with user, action, entity, tenant, and timestamp.", "#DC2626"),
    ]
    for box, title, body, accent in nodes:
        label_box(draw, box, title, body, WHITE, accent)
    arrow(draw, (395, 295), (510, 295), BLUE)
    arrow(draw, (810, 295), (925, 295), TEAL)
    arrow(draw, (1225, 295), (1340, 295), "#7C3AED")
    arrow(draw, (1490, 360), (1080, 570), "#0EA5E9")
    arrow(draw, (925, 635), (810, 635), "#DC2626")
    arrow(draw, (660, 570), (1490, 360), "#94A3B8", width=3)
    label_box(draw, (230, 830, 1570, 965), "Security Outcome", "Every protected request is allowed only when the user is authenticated, has the required permission, and is operating within the permitted tenant/geographic scope. Denied requests return explicit errors and should not mutate operational data.", MINT, TEAL)
    return save(img, "04-rbac-security-light.png")


def gis_population():
    img, draw = canvas("GIS and Population Intelligence Architecture", "How geospatial evidence supports microplanning decisions")
    boxes = [
        ((80, 200, 380, 345), "Input Layers", "Admin boundaries, facilities, settlements, population grids, POIs, roads, and custom layers.", BLUE),
        ((500, 200, 800, 345), "GIS Processing", "Containment, nearest facility, route/distance, catchment overlap, HTR and gap analysis.", TEAL),
        ((920, 200, 1220, 345), "Planning Evidence", "Catchment populations, unserved places, session reach, ETTT, and denominator confidence.", "#F97316"),
        ((1340, 200, 1640, 345), "Map UI", "Layer controls, point intelligence, routes, exports, and review panels.", "#7C3AED"),
        ((500, 555, 800, 700), "Data Model", "facilities, villages, settlements_master, population_grids, admin_boundaries, facility_catchments.", "#0EA5E9"),
        ((920, 555, 1220, 700), "Microplanning", "Facility catchments, outreach sessions, logistics, supervision, and reporting inherit GIS evidence.", "#DC2626"),
    ]
    for box, title, body, accent in boxes:
        label_box(draw, box, title, body, WHITE, accent)
    arrow(draw, (380, 272), (500, 272), BLUE)
    arrow(draw, (800, 272), (920, 272), TEAL)
    arrow(draw, (1220, 272), (1340, 272), "#F97316")
    arrow(draw, (650, 345), (650, 555), "#0EA5E9")
    arrow(draw, (1070, 345), (1070, 555), "#DC2626")
    arrow(draw, (800, 625), (920, 625), "#64748B")
    draw.text((90, 870), "Boundary rule", fill=NAVY, font=FONT_HEADER)
    text_lines(draw, "Country and tenant boundaries must constrain country-specific layers so unserved places, settlements, and catchment evidence do not bleed into neighbouring countries.", 90, 905, 1560, FONT_BODY, fill=GRAY)
    return save(img, "05-gis-population-light.png")


def data_domain():
    img, draw = canvas("Enterprise Data Model Domain Map", "Major schema areas and their operational relationships")
    center = (700, 430, 1100, 590)
    rounded(draw, center, fill=WHITE, outline=BLUE, radius=24, width=4)
    draw.text((770, 462), "Tenant-Scoped", fill=NAVY, font=FONT_TITLE)
    draw.text((765, 505), "Operational Data Model", fill=GRAY, font=FONT_HEADER)
    domains = [
        ((80, 170, 440, 310), "Identity and Access", "tenants, users, roles, permissions, device tokens, audit logs", BLUE),
        ((80, 430, 440, 570), "Administrative Geography", "regions, provinces, districts, llgs, boundaries", TEAL),
        ((80, 690, 440, 830), "Facilities and Communities", "facilities, villages, staff, CHVs, cold chain, catchments", "#F97316"),
        ((1360, 170, 1720, 310), "Microplanning", "microplans, session plans, budgets, vaccines, mobilization, approvals", "#7C3AED"),
        ((1360, 430, 1720, 570), "GIS and Population", "settlements, population grids, custom layers, HTR, recommendations", "#0EA5E9"),
        ((1360, 690, 1720, 830), "Programme Operations", "clients, vaccinations, stock, reports, supervision, VPD surveillance", "#DC2626"),
    ]
    for box, title, body, accent in domains:
        label_box(draw, box, title, body, WHITE, accent)
        if box[0] < 600:
            arrow(draw, (box[2], (box[1] + box[3]) // 2), (center[0], (box[1] + box[3]) // 2), accent)
        else:
            arrow(draw, (center[2], (box[1] + box[3]) // 2), (box[0], (box[1] + box[3]) // 2), accent)
    draw.text((90, 980), "Design intent: link geography, facility operations, microplans, stock, supervision, and surveillance through shared tenant and location keys.", fill=GRAY, font=FONT_BODY)
    return save(img, "06-data-domain-map-light.png")


def workflow_architecture():
    img, draw = canvas("Microplanning Workflow Architecture", "From national guidance to facility execution and review")
    steps = [
        ((80, 250, 330, 410), "National Plan", "Coverage targets, budget envelope, standards, campaigns, and guidance.", BLUE),
        ((430, 250, 680, 410), "Facility Baseline", "Catchment, communities, population, staff, cold chain, and readiness.", TEAL),
        ((780, 250, 1030, 410), "Wizard Steps", "Denominators, risks, sessions, staffing, vaccines, logistics, budget.", "#F97316"),
        ((1130, 250, 1380, 410), "Approval Workflow", "Facility submission, district/province/national review, lock or return.", "#7C3AED"),
        ((1480, 250, 1730, 410), "Execution", "Sessions, outreach, stock transactions, supervision, monitoring.", "#0EA5E9"),
        ((780, 650, 1030, 810), "Review and Reports", "Dashboards, missed communities, dropout, VPD, PCE, and lessons learned.", "#DC2626"),
    ]
    for box, title, body, accent in steps:
        label_box(draw, box, title, body, WHITE, accent)
    for i in range(4):
        arrow(draw, (steps[i][0][2], 330), (steps[i + 1][0][0], 330), steps[i][3])
    arrow(draw, (1605, 410), (905, 650), "#0EA5E9")
    arrow(draw, (780, 730), (205, 410), "#94A3B8", width=3)
    label_box(draw, (340, 900, 1460, 1015), "Continuous Improvement Loop", "Review outputs improve the next planning cycle by updating denominators, catchments, session strategy, stock assumptions, and supervision priorities.", MINT, TEAL)
    return save(img, "07-microplanning-workflow-light.png")


def generate_diagrams():
    return [
        ("System Context Architecture", system_context(), "System context showing users, partners, external systems, and VaxPlan platform boundaries."),
        ("Runtime and Deployment Architecture", runtime_architecture(), "Runtime and deployment view showing browser/PWA, API, services, data store, jobs, files, and integrations."),
        ("Offline-First PWA Synchronization Architecture", offline_sync(), "Offline-first synchronization flow showing baseline load, local work, outbox replay, conflict review, and cache refresh."),
        ("Authentication, RBAC, and Geographic Scope Architecture", rbac_security(), "Authentication, role-based access control, permission resolution, geographic scope, and audit trail flow."),
        ("GIS and Population Intelligence Architecture", gis_population(), "GIS and population intelligence architecture connecting boundary, settlement, route, and planning evidence."),
        ("Enterprise Data Model Domain Architecture", data_domain(), "Enterprise data model domain map grouping the main VaxPlan schema areas and shared tenant/location keys."),
        ("Microplanning Workflow Architecture", workflow_architecture(), "Microplanning workflow architecture from national plan through facility execution and review."),
    ]


def style_doc(doc: Document) -> None:
    for style_name in ["Normal"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(90, 100, 115)
    run.font.size = Pt(9)


def embed_diagrams(paths_and_captions):
    doc = Document(SOURCE_DOC)
    style_doc(doc)
    doc.add_page_break()
    title = doc.add_heading("Appendix D: VaxPlan Architecture Diagrams", level=1)
    title.runs[0].font.color.rgb = RGBColor(11, 31, 58)
    intro = doc.add_paragraph(
        "This appendix embeds light-theme PNG diagrams that summarize the VaxPlan technical architecture for developers, GIS specialists, system architects, implementation partners, and Ministry of Health technical teams. The diagrams are rendered as PNG images so they remain stable when the document is opened in Word, Google Docs, or PDF viewers."
    )
    intro.paragraph_format.space_after = Pt(8)

    for idx, (title, path, caption) in enumerate(paths_and_captions, start=1):
        if idx > 1:
            doc.add_page_break()
        heading = doc.add_heading(f"Diagram {idx}. {title}", level=2)
        heading.runs[0].font.color.rgb = RGBColor(17, 134, 178)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.7))
        add_caption(doc, caption)

    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC


def main():
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(SOURCE_DOC)
    paths = generate_diagrams()
    out = embed_diagrams(paths)
    print(out)
    print(out.stat().st_size)
    for _, path, _ in paths:
        print(path, path.stat().st_size)


if __name__ == "__main__":
    main()




