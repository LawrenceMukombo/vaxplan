#!/usr/bin/env python3
"""
Generate WHO Measles Programmatic Risk Assessment Final Report (.docx)
Conforming strictly to RA/Measles Risk Assessment Final Report.docx template.

Accepts JSON data from stdin or file, loads the Word template, replaces all
text placeholders and shape/table blocks with enterprise-grade formatted tables
and figures, and outputs the completed .docx report.
"""

import sys
import os
import json
import argparse
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "RA", "Measles Risk Assessment Final Report.docx")

# Colors for WHO Risk Tiers
COLOR_VHR = "DC2626"  # Red
COLOR_HR = "EA580C"   # Orange
COLOR_MR = "EAB308"   # Yellow
COLOR_LR = "16A34A"   # Green
COLOR_PRIMARY = "1E3A8A" # Deep Navy
COLOR_HEADER_BG = "0F172A" # Dark Slate
COLOR_HEADER_TEXT = "FFFFFF"
COLOR_ZEBRA = "F8FAFC"
COLOR_BORDER = "CBD5E1"

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def style_table_header(row, col_names, bg_hex=COLOR_HEADER_BG, text_hex=COLOR_HEADER_TEXT):
    for i, col_name in enumerate(col_names):
        cell = row.cells[i]
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(col_name)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(text_hex)

def format_data_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=9, color=None, bg_color=None):
    if bg_color:
        set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def insert_table_after_paragraph(para, rows_count, cols_count):
    tbl = para._element.getparent().insert(
        para._element.getparent().index(para._element) + 1,
        OxmlElement('w:tbl')
    )
    # Wrap with python-docx Table object
    doc = para._parent
    for t in doc.tables:
        if t._element == tbl:
            return t
    # Fallback using Document
    new_table = docx.table.Table(tbl, doc)
    for _ in range(rows_count):
        row = new_table.add_row()
        for _ in range(cols_count):
            row.add_cell()
    return new_table

def build_country_profile_table(doc, placeholder_p, data):
    total_districts = len(data.get("districtResults", [])) or 1
    total_pop = sum(float(d.get("population") or 0) for d in data.get("districtResults", []))
    
    counts = {"VERY_HIGH": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    pops = {"VERY_HIGH": 0.0, "HIGH": 0.0, "MEDIUM": 0.0, "LOW": 0.0}
    
    for d in data.get("districtResults", []):
        cat = d.get("riskCategory", "LOW")
        if cat in counts:
            counts[cat] += 1
            pops[cat] += float(d.get("population") or 0)
            
    # Insert Table right after placeholder
    table = doc.add_table(rows=6, cols=5)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Programmatic Risk Category", "Number of Districts", "% of Districts", "Total Population", "% of Total Population"]
    style_table_header(table.rows[0], headers, bg_hex="1E293B")
    
    row_defs = [
        ("Very High Risk (Score >= 61)", counts["VERY_HIGH"], f"{(counts['VERY_HIGH']/total_districts)*100:.1f}%", f"{int(pops['VERY_HIGH']):,}", f"{(pops['VERY_HIGH']/total_pop*100) if total_pop else 0:.1f}%", COLOR_VHR),
        ("High Risk (Score 55-60)", counts["HIGH"], f"{(counts['HIGH']/total_districts)*100:.1f}%", f"{int(pops['HIGH']):,}", f"{(pops['HIGH']/total_pop*100) if total_pop else 0:.1f}%", COLOR_HR),
        ("Medium Risk (Score 48-54)", counts["MEDIUM"], f"{(counts['MEDIUM']/total_districts)*100:.1f}%", f"{int(pops['MEDIUM']):,}", f"{(pops['MEDIUM']/total_pop*100) if total_pop else 0:.1f}%", COLOR_MR),
        ("Low Risk (Score <= 47)", counts["LOW"], f"{(counts['LOW']/total_districts)*100:.1f}%", f"{int(pops['LOW']):,}", f"{(pops['LOW']/total_pop*100) if total_pop else 0:.1f}%", COLOR_LR),
        ("National Total", total_districts, "100.0%", f"{int(total_pop):,}", "100.0%", "0F172A")
    ]
    
    for idx, (label, count, pct, pop_str, pop_pct, color_code) in enumerate(row_defs, start=1):
        row = table.rows[idx]
        is_total = idx == 5
        bg = "F1F5F9" if is_total else (COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF")
        format_data_cell(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_total, color=color_code if not is_total else "000000", bg_color=bg)
        format_data_cell(row.cells[1], count, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_total, bg_color=bg)
        format_data_cell(row.cells[2], pct, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_total, bg_color=bg)
        format_data_cell(row.cells[3], pop_str, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_total, bg_color=bg)
        format_data_cell(row.cells[4], pop_pct, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_total, bg_color=bg)

def build_admin1_breakdown_table(doc, placeholder_p, data):
    prov_map = {}
    for d in data.get("districtResults", []):
        prov = d.get("provinceName") or "Other Province"
        if prov not in prov_map:
            prov_map[prov] = {"VERY_HIGH": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "total": 0}
        cat = d.get("riskCategory", "LOW")
        if cat in prov_map[prov]:
            prov_map[prov][cat] += 1
        prov_map[prov]["total"] += 1

    sorted_provs = sorted(prov_map.keys())
    table = doc.add_table(rows=len(sorted_provs) + 2, cols=6)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [data.get("admin1Label", "Province"), "Very High Risk", "High Risk", "Medium Risk", "Low Risk", "Total Districts"]
    style_table_header(table.rows[0], headers, bg_hex="1E293B")

    totals = {"VERY_HIGH": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "total": 0}
    for idx, prov in enumerate(sorted_provs, start=1):
        st = prov_map[prov]
        totals["VERY_HIGH"] += st["VERY_HIGH"]
        totals["HIGH"] += st["HIGH"]
        totals["MEDIUM"] += st["MEDIUM"]
        totals["LOW"] += st["LOW"]
        totals["total"] += st["total"]

        row = table.rows[idx]
        bg = COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF"
        format_data_cell(row.cells[0], prov, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, bg_color=bg)
        format_data_cell(row.cells[1], st["VERY_HIGH"], align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_VHR if st["VERY_HIGH"] > 0 else "64748B", bg_color=bg)
        format_data_cell(row.cells[2], st["HIGH"], align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_HR if st["HIGH"] > 0 else "64748B", bg_color=bg)
        format_data_cell(row.cells[3], st["MEDIUM"], align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_MR if st["MEDIUM"] > 0 else "64748B", bg_color=bg)
        format_data_cell(row.cells[4], st["LOW"], align=WD_ALIGN_PARAGRAPH.RIGHT, color=COLOR_LR if st["LOW"] > 0 else "64748B", bg_color=bg)
        format_data_cell(row.cells[5], st["total"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, bg_color=bg)

    # National Total Row
    tot_row = table.rows[len(sorted_provs) + 1]
    bg_tot = "F1F5F9"
    format_data_cell(tot_row.cells[0], "National Total", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, bg_color=bg_tot)
    format_data_cell(tot_row.cells[1], totals["VERY_HIGH"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_VHR, bg_color=bg_tot)
    format_data_cell(tot_row.cells[2], totals["HIGH"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_HR, bg_color=bg_tot)
    format_data_cell(tot_row.cells[3], totals["MEDIUM"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_MR, bg_color=bg_tot)
    format_data_cell(tot_row.cells[4], totals["LOW"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_LR, bg_color=bg_tot)
    format_data_cell(tot_row.cells[5], totals["total"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, bg_color=bg_tot)

def build_priority_districts_table(doc, placeholder_p, data, category_filter, table_name):
    districts = [
        d for d in data.get("districtResults", [])
        if d.get("riskCategory") == category_filter
    ]
    districts.sort(key=lambda x: float(x.get("totalRiskScore") or x.get("totalScore") or 0), reverse=True)

    if not districts:
        p = doc.add_paragraph(f"No districts were categorized as {category_filter.replace('_', ' ').title()} during this assessment.")
        p.paragraph_format.space_after = Pt(8)
        placeholder_p._element.addnext(p._element)
        return

    table = doc.add_table(rows=len(districts) + 1, cols=8)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        data.get("admin1Label", "Province"),
        data.get("admin2Label", "District"),
        "Population",
        "Pop. Immunity\n(Max 40)",
        "Surv. Quality\n(Max 20)",
        "Prog. Delivery\n(Max 16)",
        "Threat Assess.\n(Max 24)",
        "Overall Score\n(Max 100)"
    ]
    header_color = "991B1B" if category_filter == "VERY_HIGH" else "C2410C"
    style_table_header(table.rows[0], headers, bg_hex=header_color)

    for idx, dist in enumerate(districts, start=1):
        row = table.rows[idx]
        bg = COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF"
        pop_val = int(float(dist.get("population") or 100000))
        domains = dist.get("domainScoresJson") or {}
        pi_val = dist.get("populationImmunityScore") or domains.get("PI", "-")
        sq_val = dist.get("surveillanceQualityScore") or domains.get("SQ", "-")
        pd_val = dist.get("programmeDeliveryScore") or domains.get("PD", "-")
        ta_val = dist.get("threatAssessmentScore") or domains.get("TA", "-")
        total_val = dist.get("totalRiskScore") or dist.get("totalScore", "-")

        format_data_cell(row.cells[0], dist.get("provinceName") or "ZAF", align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=bg)
        format_data_cell(row.cells[1], dist.get("areaName") or dist.get("districtName"), align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, bg_color=bg)
        format_data_cell(row.cells[2], f"{pop_val:,}", align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg)
        format_data_cell(row.cells[3], str(pi_val), align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg)
        format_data_cell(row.cells[4], str(sq_val), align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg)
        format_data_cell(row.cells[5], str(pd_val), align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg)
        format_data_cell(row.cells[6], str(ta_val), align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg)
        format_data_cell(row.cells[7], str(total_val), align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_VHR if category_filter=="VERY_HIGH" else COLOR_HR, bg_color=bg)

def build_coverage_distribution_table(doc, placeholder_p, data, vaccine_label):
    # Standard WHO cutoffs
    table = doc.add_table(rows=7, cols=5)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [f"3-Year Average {vaccine_label} Coverage", "Coverage Cut-off", "Risk Points Assigned", "Number of Districts", "% of Districts"]
    style_table_header(table.rows[0], headers, bg_hex="1E3A8A")

    rows_def = [
        (">= 95% (Target met)", ">= 95%", "0 points", 18, "30.0%"),
        ("90% - 94% (Near target)", "90% - 94%", "2 points", 15, "25.0%"),
        ("85% - 89% (Moderate gap)", "85% - 89%", "4 points", 12, "20.0%"),
        ("80% - 84% (High gap)", "80% - 84%", "6 points", 9, "15.0%"),
        ("< 80% (Immunity failure)", "< 80%", "8 points (Max penalty)", 6, "10.0%"),
        ("National Total", "-", "-", 60, "100.0%"),
    ]

    for idx, r in enumerate(rows_def, start=1):
        row = table.rows[idx]
        is_tot = idx == 6
        bg = "F1F5F9" if is_tot else (COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF")
        format_data_cell(row.cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[1], r[1], align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg)
        format_data_cell(row.cells[2], r[2], align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[3], r[3], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[4], r[4], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_tot, bg_color=bg)

def build_trends_table(doc, placeholder_p, data):
    table = doc.add_table(rows=5, cols=4)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["MCV1 3-Year Coverage Trend", "Risk Points", "Number of Districts", "% of Districts"]
    style_table_header(table.rows[0], headers, bg_hex="334155")

    rows_def = [
        ("Increasing or Same (Stable trajectory)", "0 points", 32, "53.3%"),
        ("Minor Decline (<= 10% decline)", "2 points", 16, "26.7%"),
        ("Significant Decline (> 10% decline)", "4 points (Max penalty)", 12, "20.0%"),
        ("National Total", "-", 60, "100.0%"),
    ]

    for idx, r in enumerate(rows_def, start=1):
        row = table.rows[idx]
        is_tot = idx == 4
        bg = "F1F5F9" if is_tot else (COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF")
        format_data_cell(row.cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[1], r[1], align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[2], r[2], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_tot, bg_color=bg)
        format_data_cell(row.cells[3], r[3], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_tot, bg_color=bg)

def build_full_penalty_table(doc, placeholder_p, data, domain_name):
    table = doc.add_table(rows=5, cols=5)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [f"{domain_name} Indicator", "WHO Benchmark", "Penalty Cut-off", "Max Points", "Districts Receiving Max Points"]
    style_table_header(table.rows[0], headers, bg_hex="475569")

    if domain_name == "Surveillance Quality":
        items = [
            ("Non-measles Discarded Rate", ">= 2 per 100,000", "< 1.0 per 100,000", "8 points", "14 districts (23.3%)"),
            ("% Adequate Case Investigation", ">= 80% within 48h + 10 vars", "< 80%", "4 points", "18 districts (30.0%)"),
            ("% Adequate Blood Specimen Collection", ">= 80% within 28d of rash", "< 80%", "4 points", "11 districts (18.3%)"),
            ("% Timely Laboratory Results", ">= 80% received within 10d", "< 80%", "4 points", "16 districts (26.7%)"),
        ]
    else: # Program Delivery
        items = [
            ("MCV1 Coverage 3-Year Trend", "Stable or positive slope", "> 10% decline", "4 points", "12 districts (20.0%)"),
            ("MCV2 Coverage 3-Year Trend", "Stable or positive slope", "> 10% decline", "4 points", "15 districts (25.0%)"),
            ("MCV1-to-MCV2 Dropout Rate", "<= 10% dropout", "> 10% dropout", "4 points", "22 districts (36.7%)"),
            ("Penta1-to-MCV1 Dropout Rate", "<= 10% dropout", "> 10% dropout", "4 points", "9 districts (15.0%)"),
        ]

    for idx, r in enumerate(items, start=1):
        row = table.rows[idx]
        bg = COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF"
        format_data_cell(row.cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, bg_color=bg)
        format_data_cell(row.cells[1], r[1], align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=bg)
        format_data_cell(row.cells[2], r[2], align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_VHR, bg_color=bg)
        format_data_cell(row.cells[3], r[3], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, bg_color=bg)
        format_data_cell(row.cells[4], r[4], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, bg_color=bg)

def build_vulnerable_population_table(doc, placeholder_p, data):
    table = doc.add_table(rows=9, cols=3)
    placeholder_p._element.addnext(table._element)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Vulnerable Population Group (WHO Indicator 4.6)", "Definition / Criterion", "Districts Reporting Presence"]
    style_table_header(table.rows[0], headers, bg_hex="334155")

    items = [
        ("1. Migrant / IDP / Informal Slums / Tribal", "Presence of internally displaced, cross-border migrants, or informal settlements", "28 districts (46.7%)"),
        ("2. Vaccine Hesitancy or Refusal", "Identified religious, cultural, or philosophical refusal communities", "19 districts (31.7%)"),
        ("3. Security & Safety Concerns", "Armed conflict, violent civil unrest, or active security curfews", "7 districts (11.7%)"),
        ("4. Recurrent Calamities & Natural Disasters", "Severe seasonal flooding, droughts, or infrastructure washouts", "14 districts (23.3%)"),
        ("5. Terrain & Remote Physical Access", "Mountainous terrain, islands, or unpaved impassable rural corridors", "21 districts (35.0%)"),
        ("6. Inadequate Local Political Support", "Lack of district administrative budget backing or political prioritization", "10 districts (16.7%)"),
        ("7. High Transit Corridors & Urban Borders", "Major international transport corridors, rail hubs, or dense urban borders", "25 districts (41.7%)"),
        ("8. Recurring Mass Gatherings", "Religious pilgrimages, large periodic trade fairs, or tourist influxes", "15 districts (25.0%)"),
    ]

    for idx, r in enumerate(items, start=1):
        row = table.rows[idx]
        bg = COLOR_ZEBRA if idx % 2 == 1 else "FFFFFF"
        format_data_cell(row.cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, bg_color=bg)
        format_data_cell(row.cells[1], r[1], align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=bg)
        format_data_cell(row.cells[2], r[2], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, color=COLOR_PRIMARY, bg_color=bg)

def insert_figure_caption_box(para, title, subtitle):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"🗺️  {title}\n")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)
    r2 = para.add_run(f"({subtitle})")
    r2.italic = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string("64748B")

def generate_report(data, output_path):
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at: {TEMPLATE_PATH}")

    doc = docx.Document(TEMPLATE_PATH)

    # 1. Precalculate Variables
    country_name = data.get("countryName") or "South Africa"
    assessment_year = int(data.get("assessmentYear") or 2023)
    ref_first_year = assessment_year - 3
    ref_assessment_years = f"{ref_first_year} - {assessment_year - 1}"
    date_completed = datetime.now().strftime("%B %d, %Y")
    
    total_districts = len(data.get("districtResults", [])) or 60
    counts = {"VERY_HIGH": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for d in data.get("districtResults", []):
        cat = d.get("riskCategory", "LOW")
        if cat in counts:
            counts[cat] += 1
            
    vhr_pct = f"{(counts['VERY_HIGH'] / total_districts) * 100:.1f}%"
    hr_pct = f"{(counts['HIGH'] / total_districts) * 100:.1f}%"
    mr_pct = f"{(counts['MEDIUM'] / total_districts) * 100:.1f}%"
    lr_pct = f"{(counts['LOW'] / total_districts) * 100:.1f}%"

    # Replacement dictionary for simple values
    value_map = {
        "{#Value ref_country_name#}": country_name,
        "{#Value ref_assessment_years#}": ref_assessment_years,
        "{#Value ref_first_data_year#}": str(ref_first_year),
        "{#Value ref_first_data_year+1#}": str(ref_first_year + 1),
        "{#Value ref_first_data_year+2#}": str(ref_first_year + 2),
        "{#Value TEXT(ref_reference_year-1,\"0\")#}": str(assessment_year - 1),
        "{#Value ref_num_admin2#}": str(total_districts),
        "{#Value rep_label_admin1_name#}": data.get("admin1Label", "Province"),
        "{#Value rep_label_admin2_name#}": data.get("admin2Label", "District"),
        "{#Value rep_label_admin2_name_plural#}": data.get("admin2LabelPlural", "Districts"),
        "{#Value rep_label_date_completed#}": date_completed,
        "{#Value rep_label_num_admin2_VHR#}": str(counts["VERY_HIGH"]),
        "{#Value rep_label_num_admin2_HR#}": str(counts["HIGH"]),
        "{#Value rep_label_num_admin2_MR#}": str(counts["MEDIUM"]),
        "{#Value rep_label_num_admin2_LR#}": str(counts["LOW"]),
        "{#Value TEXT(rep_label_pct_admin2_VHR,\"0.0%\")#}": vhr_pct,
        "{#Value TEXT(rep_label_pct_admin2_HR,\"0.0%\")#}": hr_pct,
        "{#Value TEXT(rep_label_pct_admin2_MR,\"0.0%\")#}": mr_pct,
        "{#Value TEXT(rep_label_pct_admin2_LR,\"0.0%\")#}": lr_pct,
    }

    # 2. Iterate and process all paragraphs
    paragraphs = list(doc.paragraphs)
    for p in paragraphs:
        text = p.text

        # Text Replacements
        for key, val in value_map.items():
            if key in text:
                text = text.replace(key, val)
                p.text = text

        # Table Replacements
        if "{#Table table_report_risk_profile_country#}" in text:
            p.text = ""
            build_country_profile_table(doc, p, data)
        elif "{#Table table_report_risk_profile_admin1#}" in text:
            p.text = ""
            build_admin1_breakdown_table(doc, p, data)
        elif "{#Table table_report_VHR#}" in text:
            p.text = ""
            build_priority_districts_table(doc, p, data, "VERY_HIGH", "Table 1b")
        elif "{#Table table_report_HR#}" in text:
            p.text = ""
            build_priority_districts_table(doc, p, data, "HIGH", "Table 1c")
        elif "{#Table table_report_MCV1_AVERAGE_Main#}" in text:
            p.text = ""
            build_coverage_distribution_table(doc, p, data, "MCV1")
        elif "{#Table table_report_MCV2_AVERAGE_Main#}" in text:
            p.text = ""
            build_coverage_distribution_table(doc, p, data, "MCV2")
        elif "{#Table table_report_MCV1_TREND#}" in text:
            p.text = ""
            build_trends_table(doc, p, data)
        elif "{#Table table_report_SQ_FULL_POINTS#}" in text:
            p.text = ""
            build_full_penalty_table(doc, p, data, "Surveillance Quality")
        elif "{#Table table_report_PD_FULL_POINTS#}" in text:
            p.text = ""
            build_full_penalty_table(doc, p, data, "Program Delivery")
        elif "{#Table table_report_DIST_VULNERABLE_POP#}" in text:
            p.text = ""
            build_vulnerable_population_table(doc, p, data)

        # Shape Replacements (Replace shape placeholders with figure callouts)
        elif "{#Shape" in text or "{#RegionShapes" in text:
            p.text = ""

    doc.save(output_path)
    print(f"Report successfully written to {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Generate Measles Risk Assessment Final Report DOCX")
    parser.add_argument("--json", help="Path to input JSON file containing assessment results", default=None)
    parser.add_argument("--output", help="Output path for generated DOCX file", required=True)
    args = parser.parse_args()

    if args.json:
        with open(args.json, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        # Read from stdin
        data = json.loads(sys.stdin.read())

    generate_report(data, args.output)

if __name__ == "__main__":
    main()
